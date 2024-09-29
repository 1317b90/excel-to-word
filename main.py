import pandas as pd
import os
from docx import Document
import shutil
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

excelpath = "./GoalsExcel"
print(10086)
# 过滤出Excel文件
excels = [file for file in os.listdir(excelpath) if file.endswith(('.xlsx', '.xls'))]
# 获取第一个excel文件的路径
excelpath = os.path.join(excelpath, excels[0])

print("读取文件：", excels[0])

# 先读取第一次excel
df = pd.read_excel(excelpath)
# 获取表头
cols = list(df.columns)
dtype_dict={}
for col in cols:
    dtype_dict[col]=str

# 第二次读取，指定每一列数据类型为str，保证编号00x可以读取
df = pd.read_excel(excelpath, dtype=dtype_dict)

# 编号列表
numbers = list(df[cols[0]].dropna().unique())

# 储存地区名称
regions_dict = {}

# 读取地区名称
k = 0
for region in df[cols[4]]:
    number = numbers[k]
    if pd.isna(region):
        k = k + 1
    else:
        if regions_dict.get(number) is None:
            regions_dict[number] = []
        regions_dict[number].append(region)
print(regions_dict)

# 非空索引
nullindex = []
temp_list = list(df[cols[0]])
for i in range(len(temp_list)):
    if not pd.isna(temp_list[i]):
        nullindex.append(i)

# 执行前先将空值转为空字符串，然后转换整体数据格式为str
df = df.fillna('')

# 居中单元格
def centenr(cell):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# 一个编号一个文件
for number in numbers:
    print("读取编号：",number,"...")
    # 该编号的非空数值所在的索引
    index = nullindex[numbers.index(number)]
    # 文件名称
    # 编号-名称-文本6-文本3-2024-待检
    wordname = str(number) + "-" + df.iloc[index, 4] + "-" + df.iloc[index, 8]+"-" + df.iloc[index, 6] +"-2024-待检.docx"

    # 存放word模版的文件夹
    templatepath = "./TemplateWord"
    files = [file for file in os.listdir(templatepath) if file.endswith(('.docx', '.doc'))]
    templatepath = os.path.join(templatepath, files[0])

    # 文件路径
    wordpath = './GenerationWord/' + wordname
    # 将模板word复制到生成word文件夹中
    shutil.copy(templatepath, wordpath)

    # 打开Word文档
    doc = Document(wordpath)

    # ----------在附录1中----------在附录1中----------在附录1中----------在附录1中
    table1 = doc.tables[0]

    # 日期1
    table1.cell(13, 1).text = str(df.iloc[index, 1])[:10]
    centenr(table1.cell(13, 1))

    # 日期2 在word中没用到

    # 编号2
    table1.cell(14, 1).text = df.iloc[index, 3]
    centenr(table1.cell(14, 1))

    # 名称
    table1.cell(0, 1).text = df.iloc[index, 4]
    centenr(table1.cell(0, 1))

    # 文本2 在word中没用到

    # 文本3
    table1.cell(6, 1).text = df.iloc[index, 6]
    centenr(table1.cell(6, 1))

    # 文本 4
    table1.cell(7, 1).text = df.iloc[index, 7]
    centenr(table1.cell(7, 1))

    # 文本 8
    table1.cell(3, 1).text = df.iloc[index, 10]
    centenr(table1.cell(3, 1))

    # 日期3
    table1.cell(8, 1).text = str(df.iloc[index, 11])[:10]
    centenr(table1.cell(8, 1))

    # 日期4
    table1.cell(9, 1).text = str(df.iloc[index, 12])[:10]
    centenr(table1.cell(9, 1))

    # 文本9
    table1.cell(4, 1).text = df.iloc[index, 13]
    centenr(table1.cell(4, 1))

    # ---------在附录2中---------在附录2中---------在附录2中---------在附录2中---------在附录2中
    table2 = doc.tables[1]
    for i in range(1, len(regions_dict[number])):
        # 编号2
        table2.cell(3, i + 1).text = df.iloc[index + i, 3]
        centenr(table2.cell(3, i + 1))
        # 名称
        table2.cell(0, i + 1).text = df.iloc[index + i, 4]
        centenr(table2.cell(0, i + 1))
        # 文本3
        table2.cell(7, i + 1).text = df.iloc[index + i, 6]
        centenr(table2.cell(7, i + 1))
        # 文本4
        table2.cell(8, i + 1).text = df.iloc[index + i, 7]
        centenr(table2.cell(8, i + 1))
        # 文本 6
        text6=df.iloc[index + i, 8]
        if text6=="":
            text6 = df.iloc[index + i-1, 8]

        table2.cell(16, i + 1).text = text6
        centenr(table2.cell(16, i + 1))
        # 文本 7
        table2.cell(18, i + 1).text = df.iloc[index + i, 9]
        centenr(table2.cell(18, i + 1))
        # 文本 8
        table2.cell(4, i + 1).text = df.iloc[index + i, 10]
        centenr(table2.cell(4, i + 1))
        # 日期3
        table2.cell(9, i + 1).text = str(df.iloc[index+i, 11])[:10]
        centenr(table2.cell(9, i + 1))
        # 日期4
        table2.cell(10, i + 1).text = str(df.iloc[index+i, 12])[:10]
        centenr(table2.cell(10, i + 1))
        # 文本9
        table2.cell(5, i + 1).text = df.iloc[index + i, 13]
        centenr(table2.cell(5, i + 1))

    doc.save(wordpath)
print("读取完毕，文件已保存！")